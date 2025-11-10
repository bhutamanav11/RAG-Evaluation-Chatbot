import fitz  # PyMuPDF
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from typing import List, Dict, Any, Optional
import os
import re

class DocumentProcessor:
    def __init__(self, chunk_size: int = 250, chunk_overlap: int = 40):
        """
        IMPROVED: Smaller chunks (250 tokens ~= 200-300 words) with 15% overlap
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Use semantic separators - respect document structure
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",  # Multiple blank lines (section breaks)
                "\n\n",    # Paragraph breaks
                "\n",      # Line breaks
                ". ",      # Sentence breaks
                "! ",      # Exclamation sentences
                "? ",      # Question sentences
                "; ",      # Semi-colons
                ", ",      # Commas
                " ",       # Spaces
                ""         # Characters
            ],
            is_separator_regex=False
        )
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text - remove boilerplate, headers, navigation
        """
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove common boilerplate patterns
        boilerplate_patterns = [
            r'Page \d+ of \d+',
            r'©.*?All rights reserved',
            r'Copyright.*?\d{4}',
            r'Table of Contents',
            r'^\d+\s*$',  # Page numbers on their own line
            r'[A-Z\s]{10,}',  # Long ALL CAPS headers
            r'www\.\S+',  # URLs
            r'https?://\S+',  # URLs
            r'\[.*?\]',  # Square bracket annotations
            r'^\s*\d+\.\s*$',  # Lone numbering
        ]
        
        for pattern in boilerplate_patterns:
            text = re.sub(pattern, '', text, flags=re.MULTILINE | re.IGNORECASE)
        
        # Remove navigation text patterns
        nav_patterns = [
            r'(Home|About|Contact|Privacy Policy|Terms of Service)',
            r'(Click here|Learn more|Read more)',
            r'(Next|Previous|Back to top)',
        ]
        
        for pattern in nav_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        # Clean up residual whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = text.strip()
        
        return text
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with improved quality"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num, page in enumerate(doc):
                # Extract text with layout preservation
                page_text = page.get_text("text")
                
                # Skip if page is mostly empty (< 50 chars)
                if len(page_text.strip()) < 50:
                    continue
                
                text += page_text + "\n\n"
            
            doc.close()
            
            # Clean the extracted text
            text = self.clean_text(text)
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX with improved quality"""
        try:
            doc = Document(file_path)
            text = ""
            
            for para in doc.paragraphs:
                para_text = para.text.strip()
                
                # Skip empty paragraphs
                if not para_text:
                    continue
                
                # Skip very short paragraphs that are likely headers/page numbers
                if len(para_text) < 10:
                    continue
                
                text += para_text + "\n\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            # Clean the extracted text
            text = self.clean_text(text)
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def semantic_chunk(self, text: str) -> List[str]:
        """
        Create semantic chunks that respect document structure
        """
        # First split by the recursive splitter
        base_chunks = self.text_splitter.split_text(text)
        
        # Post-process chunks for quality
        semantic_chunks = []
        
        for chunk in base_chunks:
            chunk = chunk.strip()
            
            # Skip very short chunks (likely artifacts)
            if len(chunk) < 30:
                continue
            
            # Skip chunks that are just numbers or dates
            if re.match(r'^[\d\s\-/.,]+$', chunk):
                continue
            
            # Ensure chunk ends at sentence boundary if possible
            if not chunk.endswith(('.', '!', '?', '"', "'")):
                # Try to find last sentence ending
                last_sentence = max(
                    chunk.rfind('.'),
                    chunk.rfind('!'),
                    chunk.rfind('?')
                )
                if last_sentence > len(chunk) * 0.7:  # If found in last 30%
                    chunk = chunk[:last_sentence + 1]
            
            semantic_chunks.append(chunk)
        
        return semantic_chunks
    
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process document with improved chunking"""
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            raw_text = self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            raw_text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Create semantic chunks
        chunks = self.semantic_chunk(raw_text)
        
        # Create document structure with metadata
        document_chunks = []
        for i, chunk in enumerate(chunks):
            document_chunks.append({
                'content': chunk,
                'filename': filename,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'file_type': file_extension,
                'chunk_length': len(chunk)
            })
        
        return {
            'filename': filename,
            'file_type': file_extension,
            'raw_text': raw_text,
            'chunks': document_chunks,
            'chunk_count': len(chunks)
        }
    
    def adaptive_chunking(self, text: str, conversation_turn: int, 
                         model_name: str) -> List[str]:
        """
        Adaptive chunking based on conversation context
        Smaller chunks for early turns, slightly larger for later turns
        """
        # Base chunk size varies with conversation length
        if conversation_turn <= 5:
            chunk_size = 250  # Small chunks for precision
        elif conversation_turn <= 15:
            chunk_size = 300  # Medium chunks
        else:
            chunk_size = 350  # Slightly larger for context
        
        # Model-specific adjustments
        if 'gpt-4' in model_name:
            chunk_size = int(chunk_size * 1.1)
        elif 'claude' in model_name:
            chunk_size = int(chunk_size * 1.05)
        
        # Create adaptive splitter
        adaptive_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=int(chunk_size * 0.15),  # 15% overlap
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        return adaptive_splitter.split_text(text)